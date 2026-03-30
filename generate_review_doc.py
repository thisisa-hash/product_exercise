from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('Product Manager Assessment — Candidate Reviews', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Prepared by: AI-assisted evaluation | Based on calibrated framework')
doc.add_paragraph('')

# ─── JERT ───
doc.add_heading('JERT', 1)

doc.add_heading('Exercise 1 — Offers Page', 2)
for b in [
    'Funnel is Entry → View → Click → Activate → Deposit with 4 illustrative rows. No mention of Gift icon, offer sections (Sports/Casino/All), banners or any element specific to the Offers Page as a product.',
    'Only 3 KPIs defined (CTR, Activation Rate, Deposit Rate) — all derived from the funnel steps themselves. No bounce rate, scroll depth, repeat visit rate, or anything that reflects how users browse and discover offers.',
    'Success defined purely as deposit conversion — misses the hub/discovery role of the product entirely.',
    'Additional KPIs section offers only 3 items (conversion by entry point, time from activation to deposit, repeat usage) qualified with "if I had access to more data" — a PM should propose these unconditionally.',
    'No engagement with the Offers Page screenshots file at all.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 2 — Missions', 2)
for b in [
    'Used 3 cherry-picked rows (Apr 6, Apr 20, May 4) from 13 weeks of available data — no trend, no monthly view, no aggregation.',
    'Never mentioned public vs private missions once — the core personalization lever in the product and the most actionable variable in the dataset.',
    'Sport categories (football, basketball, tennis, mixed) completely ignored — the Category tab was not touched.',
    '"Product improvements" are: simplify missions, improve visibility, optimize difficulty — generic enough to apply to any gamification product.',
    '"Drop the product" conclusion is correct but based on zero data: "high opt-ins indicate product-market fit" — that is one number from one tab.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 3 — Loyalty', 2)
for b in [
    'Correctly understood there is no opt-in.',
    'Framed around three dimensions: Adoption & Engagement, Progression Dynamics, Behavioral Impact — reasonable structure.',
    'Provided illustrative funnel (100K → 72K → 28K → 9.5K) and a cohort retention example (week 4: 48%, VIP week 4: 70%+) — shows effort to make it tangible.',
    'Metrics proposed are generic — no reward cost vs GGR, no coin redemption rate, no churn by stage. Stops at "are users progressing?" without asking "is this profitable for the business?"',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 4 — NSM', 2)
for b in [
    'Chose Sportsbook Missions, NSM: Average Completed Missions per Active User — defensible but essentially a repackaging of completion rate.',
    'Supporting metrics are well-chosen (opt-in rate, completion rate, avg bets per user, retention of mission users) — the best part of this exercise.',
    'Prioritization framework (impact x effort x confidence) is textbook correct but generic — no connection to Novibet\'s specific product context.',
    'Never argued why Missions over Offers Page or Loyalty — the choice feels default, not strategic.',
]:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Final call: No Hire.  ').bold = True
p.add_run('The submission is well-structured but hollow. Every exercise describes methodology instead of demonstrating product judgment. The Missions analysis — where real data was available — is the clearest disqualifier: 3 rows used, two of the three key variables ignored, improvements that could apply to any product. No evidence of a PM who asks "what is this product trying to do for the user and for the business?"')

p2 = doc.add_paragraph()
p2.add_run('AI signal: High.  ').bold = True
p2.add_run('Identical template structure across all 4 exercises, pervasive "I would" distancing language, no data engagement despite having it, no supporting file submitted, no voice or opinion anywhere.')

doc.add_paragraph('')

# ─── GUILLERMO ───
doc.add_heading('GUILLERMO', 1)

doc.add_heading('Exercise 1 — Offers Page', 2)
for b in [
    'Defined three distinct funnels reflecting how users actually experience the product differently: Discovery (how they find it), Navigation (how they browse inside it), Activation & Completion (how they act on it) — shows he understood the Offers Page as a hub, not a landing page.',
    'Events properly categorized into Acquisition, Engagement and Conversion — with specific events like "Offer Category Filtered/Tapped" and "Offer Abandoned" that show he thought about the browsing experience.',
    'Success defined as "a growing share of active players discover and activate at least one promotion per session, increasing bet frequency and measurable GGR" — connects user action directly to business outcome.',
    'Additional KPIs go beyond basics: Offer to Bet Ratio, Cross-product Activation (casino players activating sports offers and vice versa), Category Heatmap CTR — shows strategic product thinking about the hub\'s role across verticals.',
    'Weak spot: illustrative report is only 4 headline numbers. The Brief acknowledges this was intentional ("measurement architecture over mock data") but it leaves the report thin.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 2 — Missions', 2)
for b in [
    'Identified the June completion drop precisely: active users completing at least one mission fell from 2,499 (May 18) to 1,234 (Jun 29) — a 51% decline — while opt-in users remained stable at ~4,600–4,750. The problem is not discovery, it\'s completion.',
    'Correctly identified private missions as a personalization lever and flagged they were underused — only 0–5 per week across 13 weeks — with a clear recommendation to expand targeting for high-value segments.',
    'Sport category analysis is strong: football as backbone, basketball disappearing in June (end of NBA), Mixed missions peaking on the best week (May 18: 6 Mixed, 25,432 opt-ins), tennis unable to fill the gap. Connected seasonality to supply-side decisions.',
    'Success definition includes concrete thresholds: at least 20% of weekly active users opt-in, CR stays above 35%, no widening gap between opt-in and completion users.',
    '"Drop the product" response is structured around four distinct arguments with data: absolute scale (182K opt-ins), growing user base (+30%), supply-side root cause, and the product\'s unique retention mechanic (rewards regardless of win/loss).',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 3 — Loyalty', 2)
for b in [
    'Correctly understood no opt-in, correctly distinguished Premium (one-off) vs VIP (monthly recurring).',
    'Time-phased structure — Launch (Jan–Feb), Growth (Mar–May), Maturity (Jun–Aug) plus a pre-launch baseline (Oct–Dec 2024) — is sharp product thinking: you cannot assess a program without a before/after reference.',
    'Illustrative dashboard is the most business-grounded: 68% Premium completion, 31% Premium to VIP, +22% GGR uplift vs non-loyalty cohort, 4.7x sessions for VIP users — these numbers tell a product story.',
    'Additional KPIs include Reward to Cost Ratio and GGR Lift (loyalty vs control cohort) — demonstrates awareness that loyalty programs need to be commercially justified, not just engagement-tracked.',
    'Brief adds a sharp further exploration question: "Are players earning coins but not redeeming them? Low redemption would indicate an awareness or UX issue in the reward flow, not a product design problem" — this is PM thinking.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 4 — NSM', 2)
for b in [
    'Chose Offers Page — Weekly Active Players with at least 1 Offer Activation — and argued why this metric filters out passive browsing while capturing real intent: a player who activates has received the core value of the product.',
    'The decision framework is the most actionable of all candidates: Awareness Driver (Entry Point CTR) → Engagement Driver (Offer Click Rate) → Conversion Driver (Opt-In Rate), each with a specific initiative attached. This is how a PM actually runs a weekly product review.',
    'Explicitly noted how to detect trade-offs: "more offers shown but lower CTR may signal relevance issues rather than a volume problem" — shows he has thought about how the metric can mislead.',
    'Weak spot: did not argue why Offers Page over Loyalty or Missions. The choice is well-justified on its own terms but the comparative reasoning is absent.',
]:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Final call: Hire — #1.  ').bold = True
p.add_run('Guillermo demonstrates consistent product-first thinking across all four exercises. He understands what each product is trying to do for the user, connects observations to business outcomes, and proposes improvements that are strategic rather than generic. The Brief in particular shows a candidate who can reflect on his own methodology — a strong PM signal.')

p2 = doc.add_paragraph()
p2.add_run('AI signal: Low.  ').bold = True
p2.add_run('Has a distinct voice, genuine data observations grounded in the actual dataset, personal assumptions in the Brief that feel authored, and a visual presentation style that reflects individual choices.')

doc.add_paragraph('')

# ─── GALLARDO ───
doc.add_heading('GALLARDO', 1)

doc.add_heading('Exercise 1 — Offers Page', 2)
for b in [
    'The work document is the most comprehensive event taxonomy of all candidates — Traffic & Discovery, Engagement & Interaction, Offer Discovery & Selection, Conversion & Funnel, Retention & Behavioral — mapped to a full Buyer Journey (Awareness → Consideration → Conversion → Retention). Nothing is missing.',
    'Illustrative report is the richest: 120K users, entry points (Gift icon 48%, homepage banner 27%, push/email 15%), scroll depth split (22%/46%/32%), section click distribution, per-offer CTR and activations, full funnel (120K→54K→26K→18.5K), and post-offer behavior (loyalty users returning at 68% vs 39% non-loyalty).',
    'The 68% vs 39% retention comparison is a genuinely strong product insight — it demonstrates why the Offers Page has business value beyond just activations.',
    'Additional KPIs include "time to activation" and "steps to activation" as friction indicators — more actionable than what other candidates proposed.',
    'Weak spot: the final PDF condenses all of exercise 1 into a single short paragraph. The depth lives in the work document — a PM presenting to leadership would need to bridge that gap.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 2 — Missions', 2)
for b in [
    'Pinpointed the exact inflection point: May 25 (45.47% CR, 40% Private) → June 1 (26.70%, 0 Private, 5 missions) — framed as a signal of an operational or configuration change, not a product failure. That is precise PM thinking.',
    'Quantified the opt-in/completion divergence explicitly: opt-in rate grew from ~16% to over 17% while completion rate fell from ~8–9% to 4.5% by June 29. Correctly diagnosed as a completion-stage issue, not a discovery problem.',
    'Full Category analysis: April Football+Basketball dominant, May most diverse with Mixed missions (43% of May 18 week), June collapsed to Football+Tennis only — identified as likely a seasonality or operational constraint and connected directly to the performance drop.',
    'Private missions framed as a personalization strategy: recommended a 60–70% Public / 30–40% Private mix as optimal, with A/B tests to validate.',
    '"Drop the product" response is the most structured: four specific reasons (strong absolute scale, growing opt-in user base, supply-side root cause, unique retention mechanic) — mirrors how a PM would defend a product in a leadership meeting.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 3 — Loyalty', 2)
for b in [
    'Correctly understood no opt-in from the first line, correctly distinguished Premium vs VIP mechanics.',
    'Introduced a Reward Effectiveness dimension (redemption rate, time between earning and redemption, behavior before and after milestones) that no other candidate raised — this is a product maturity question: are players actually getting value from the rewards?',
    'Perimeter is well-defined and respects the assignment guidance: casino-only, excludes competitive comparisons, excludes over-focus on coin numerics.',
    'Four-dimension framework (Adoption, Engagement & Progression, Reward Effectiveness, Retention & Behavioral Impact) covers the full lifecycle of the product.',
    'Weak spot: no visuals, no pre-launch baseline (Oct–Dec 2024) for comparison — Guillermo\'s time-phased approach with a before/after anchor is sharper here.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 4 — NSM', 2)
for b in [
    'Chose Casino Loyalty — Monthly Active Casino Loyalty Users — and articulated all four design pillars of the product in the rationale: frictionless, activity-driven, progressive, retention-focused. Shows he actually read and internalized the product description.',
    'Explicitly distinguished the NSM from vanity metrics: "unlike total registered users, this captures ongoing, value-creating usage" — a PM who understands what a metric is really measuring.',
    'Decision framework is crisp and tied to the product\'s structure: if NSM declines → investigate discovery; if Premium completion low → simplify milestones; if VIP stagnates → optimize rewards; if engagement drops → introduce progression boosters.',
    'Secondary metrics are all genuinely complementary and non-redundant: Premium completion rate, VIP monthly active users, avg coins per active user, retention lift vs non-loyalty.',
]:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Final call: Hire — #2.  ').bold = True
p.add_run('Gallardo is analytically the deepest of all candidates and that depth is mostly in service of product decisions — the private mission personalization argument, the reward effectiveness dimension in loyalty, the operational change hypothesis on May 25. The main risk is that the work document reads like an analyst report rather than a PM presentation. Worth probing in interview on how he communicates product decisions to non-analytical stakeholders.')

p2 = doc.add_paragraph()
p2.add_run('AI signal: Low-Medium.  ').bold = True
p2.add_run('The analytical granularity feels genuinely worked through. However, the writing in the work document is unusually thorough and formally structured throughout — possible heavy AI editing even if the thinking is his own.')

doc.add_paragraph('')

# ─── GUILHERME COSTA ───
doc.add_heading('GUILHERME COSTA', 1)

doc.add_heading('Exercise 1 — Offers Page', 2)
for b in [
    'Opened with a Brazil-specific regulatory context (bonus restrictions since market regulation) and used the actual screenshots to identify UX issues — absence of CTAs, offer counters showing 496/1,747 days to expiry (makes FOMO meaningless), no clear hierarchy for featured offers. No other candidate looked at the product this critically.',
    'KPIs structured into a 5-step sequential funnel (page view → banner click → CTA → register → bet placed) with drop-off hypotheses per step. Separate platform breakdown (Mobile App/Mobile Web/Desktop) shows platform-aware PM thinking.',
    'Success definition is tiered by timeframe: Short-term weekly (10% click rate, 25% opt-in of clicks, <60% exit), Medium-term monthly (10% offer-to-bet CVR, 12% GGR from offers), Long-term quarterly (multi-offer engagement, bonus cost ratio ≤25%) — most structured success framework of all candidates.',
    'Additional KPIs include Bonus Cost Ratio, CAC via offers, GGR per offer — commercially grounded.',
    'Transparently flagged that dashboard figures were produced with AI for illustration purposes — honest and appropriate.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 2 — Missions', 2)
for b in [
    'Full 13-week data used with clean summary table and monthly aggregation. Identified the key structural insight: opt-in rate kept rising in June (16.7% → 17.3%) while completions collapsed — player interest did not drop, only their ability to complete missions did. The sharpest single observation in the entire dataset.',
    'Connected sport category mix directly to performance with evidence: May best week (May 18: 14 missions, 3F/5B/6 Mixed) vs June collapse (0 Basketball, 0 Mixed). Validated the pattern with April 27 (Football+Tennis only → CR dropped from 36.9% to 31.8%).',
    'Raised a secondary issue no other candidate spotted: the User Completion Rate (completing users / opt-in users) dropped from 60.7% in April to 26% by late June — and this decline started in May before the June category shift, suggesting mission fatigue or difficulty calibration drift.',
    'Called out that football is underrepresented at 44.4% of missions despite representing 80%+ of the Brazilian sports betting market — a market-specific insight showing genuine local knowledge.',
    '"Drop the product" response is the most precise: "Missions does not have a product problem, but a content curation problem." Proposed a controlled A/B experiment (Missions On vs Off for matched cohorts over 6 weeks) as the correct decision-making tool.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 3 — Loyalty', 2)
for b in [
    'Correctly understood no opt-in. Correctly identified Premium (5 milestones) and VIP (7 milestones) structure.',
    'Opened with Brazilian market context (28M active players, 68.3% men, 51.3% aged 18–40) and connected this demographic to gaming culture and dopamine-driven loyalty mechanics — strategic framing that goes beyond the product itself.',
    'Spotted a structural design risk: M3→M4 requires 25,000 coins vs M2→M3 requiring 10,000 — a 2.5x jump — and proposed an interim milestone at ~28,000–30,000 coins to reduce drop-off. No other candidate engaged at this level of product design detail.',
    'Proposed four audience segments: New players, Established players, Reactivated players, and a Control group (Aug–Dec 2024) — and flagged January 2025 results are likely outliers due to market regulation disruption (Blask Index fell 28%). Context-awareness is exceptional.',
    'Risks section is unique: includes acquisition cost, milestone drop-off, retention cost, reactivation cost, and programme cost vs other onboarding alternatives — treats the loyalty programme as a business investment, not just a product feature.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 4 — NSM', 2)
for b in [
    'Chose Sportsbook Missions — NSM: Weekly Transacting Mission Users (WTMU) — defined as unique players who both opted in AND completed at least one mission in a week. Most precisely defined NSM of all candidates.',
    'Acknowledged the distinction between metrics that capture intent (opt-ins) vs. platform decisions (number of missions) vs. actual business impact — showed conceptual clarity about what an NSM should measure.',
    'Secondary metrics include a Category Diversity Index using HHI (Herfindahl-Hirschman Index) as a concentration measure — sophisticated and directly addresses the June collapse root cause.',
    'NSM framework is the most operationally developed: diagnostic decision table, initiative prioritisation with quantified WTMU impact estimates (+30–40% for category diversity), quarterly OKR structure with targets and owners, and A/B test governance rules (95% confidence, 5% MDE, 2–4 week minimum).',
]:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Final call: Hire — #1.  ').bold = True
p.add_run('Brazilian market context, UX critique of the actual product, the mission fatigue hypothesis, the milestone gap observation in loyalty, and the WTMU definition all demonstrate a PM who genuinely engaged with the material. The transparency about AI use is a signal of good judgement — he used it as a tool, acknowledged it, and the intellectual contribution is clearly his own.')

p2 = doc.add_paragraph()
p2.add_run('AI signal: Low-Medium (disclosed).  ').bold = True
p2.add_run('He explicitly flagged AI use twice (dashboard illustrations, NSM framework steps). The analytical thinking and market context are his own. Not a red flag — it is appropriately disclosed.')

doc.add_paragraph('')

# ─── BRUNO DI MAURO ───
doc.add_heading('BRUNO DI MAURO', 1)

doc.add_heading('Exercise 1 — Offers Page', 2)
for b in [
    'Identified the cross-sell opportunity (Sportsbook customers activating casino offers and vice versa) and built two separate illustrative funnels for each direction — a genuinely practical product insight reflecting how offers pages work commercially.',
    'Proposed separating fixed offers from ad hoc offers as a UX improvement — creating a dedicated fixed-offers area vs a dynamic ad hoc area shows product structural thinking.',
    'Connected the Offers Page to CRM explicitly — checking offer performance against mailing campaigns, A/B testing offer messaging — demonstrates understanding that this product does not sit in isolation.',
    'Weak spot: events and KPIs described in a single paragraph without clear structure. No funnel table, no metric targets, no success thresholds defined.',
    'Success definition is mostly qualitative ("main success is increasing customer retention") — correct direction but not developed into measurable criteria.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 2 — Missions', 2)
for b in [
    'Identified all key pain points correctly: fewer missions in June, sudden stop of Mixed Missions, decrease in Private missions, stable opt-in rate but declining completions — shows he read the data properly.',
    'Raised a hypothesis that stands out: possibly harder betting requirements to release bet credits, causing opt-ins without completions — a product mechanics hypothesis other candidates missed.',
    'Asked the right questions for the product team: why did we cut Mixed missions? Why reduce Private? Are we targeting the right events? Why so much Tennis in June? — consultative framing is strong PM behaviour.',
    'Brazil-specific angle: explicitly mentioned the absence of opening account offers due to legislation, making Missions a must-keep product vs competitors — relevant market context.',
    'Weak spot: no data tables, no quantification, no trend analysis. Reads as a series of observations and questions rather than a structured data-driven narrative.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 3 — Loyalty', 2)
for b in [
    'Correctly understood no opt-in.',
    'Proposed a practical two-period comparison (pre vs post-launch) with specific data requests: active casino users, new sign-ups, players per tier, stakes, GGR, top 50 most played games — concrete and actionable.',
    'Raised a commercially smart idea: multiplying coin points for specific games (x1.5 or x2.0) where the operator has better provider deals, timed to low-traffic hours with on-site push notifications to Sportsbook users — a real commercial operations idea showing practical iGaming experience.',
    'Weak spot: no analytical framework (4 dimensions, progression funnel, etc.). Reads more as a list of data requests. No mention of Premium vs VIP mechanics being central to the design.',
]:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Exercise 4 — NSM', 2)
for b in [
    'Chose Sportsbook Missions — NSM: "Increase stakes/slips related to Sportsbook Missions" — the most commercially grounded NSM choice, directly tied to revenue.',
    'Visual map (New Sign-Ups → Private Missions → Active Users Converted → Missions Completed → Increase stakes/slips) shows a logical product flow.',
    'Secondary metrics are relevant: new sign-ups, customer retention, weekly private missions, optimised missions (competitive benchmarking) — the last one reflects his operator experience.',
    'Weak spot: NSM rationale is underdeveloped. "That is the main goal" is stated but not argued. No explanation of why stakes/slips is better than completion rate or opt-in rate as the north star.',
]:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run('Final call: Lean Hire.  ').bold = True
p.add_run('Bruno brings genuine iGaming operator experience that shows up in practical, commercially grounded observations — the cross-sell funnel, the coin multiplier idea, the CRM-offers connection, the Brazilian legislation context. He thinks like someone who has run these products. However, the submission lacks structure, data engagement, and analytical depth. The ideas are real but underdeveloped. Worth a conversation — the operator instincts are there, the PM framework needs strengthening.')

p2 = doc.add_paragraph()
p2.add_run('AI signal: Low.  ').bold = True
p2.add_run('Strong personal voice, practical operator-level ideas, informal and direct writing style that is clearly his own. The Excel file submitted shows genuine data work.')

doc.add_paragraph('')

# ─── SUMMARY TABLE ───
doc.add_heading('Final Summary', 1)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Candidate'
hdr[1].text = 'Final Call'
hdr[2].text = 'AI Signal'
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True

for i, (c, f, a) in enumerate([
    ('Guilherme Costa', 'Hire — #1',    'Low-Medium (disclosed)'),
    ('Guillermo',       'Hire — #2',    'Low'),
    ('Gallardo',        'Hire — #3',    'Low-Medium'),
    ('Bruno Di Mauro',  'Lean Hire',    'Low'),
    ('JERT',            'No Hire',      'High'),
]):
    row = table.rows[i+1].cells
    row[0].text = c
    row[1].text = f
    row[2].text = a

doc.save('Candidate_Reviews_PM_Assessment.docx')
print('Done')
